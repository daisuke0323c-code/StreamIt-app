# -*- coding: utf-8 -*-
import json
import os
import uuid
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional

import streamlit as st

# 取込用（存在しなければフォールバック）
try:
    import pandas as pd
except Exception:
    pd = None

try:
    from docx import Document
except Exception:
    Document = None  # python-docx が無ければプレーンテキスト抽出にフォールバック

import re
import tempfile
from pathlib import Path

try:
    import pypandoc
except Exception:
    pypandoc = None

# Web 取り込み用（存在しなければフォールバック）
try:
    import requests
except Exception:
    requests = None

try:
    import trafilatura
except Exception:
    trafilatura = None

try:
    from readability import Document as ReadabilityDocument
except Exception:
    ReadabilityDocument = None

try:
    from bs4 import BeautifulSoup
except Exception:
    BeautifulSoup = None

try:
    import html2text as _h2t
except Exception:
    _h2t = None


# ============ スタイル/CSS ============
st.markdown("""
<style>
/* トップバー（薄いカプセル） */
.agent-topbar{
  display:flex; align-items:center; justify-content:space-between; gap:8px;
  padding:6px 10px;
  border-radius: 999px;
  border: 1px solid rgba(120,140,180,.25);
  background: rgba(80,120,200,.06);
}
[data-theme="dark"] .agent-topbar{
  border-color: rgba(120,160,220,.15);
  background: rgba(80,140,240,.08);
}
.agent-topbar.is-current{
  background: rgba(90,140,240,.12);
  border-color: rgba(120,160,240,.28);
}

/* タイトルと小チップ */
.agent-title{ display:flex; align-items:center; gap:8px; font-weight:700; letter-spacing:.2px; }
.agent-chip{
  font-size: 11px; padding: 2px 8px; border-radius:999px;
  background: rgba(90,140,240,.18); color:#2e6fed; border:1px solid rgba(90,140,240,.35);
}
[data-theme="dark"] .agent-chip{ color:#9cc0ff; }

/* 現在行用の控えめチップ */
.agent-badge-current{
  font-size: 11px; padding:2px 8px; border-radius:999px;
  background: rgba(90,140,240,.14); color:#2e6fed; border:1px solid rgba(90,140,240,.28);
}

/* ミニボタン行 */
.btn-row{ display:grid; grid-template-columns: repeat(4, 1fr); gap:6px; margin-top:6px; }
.small-btn button{ padding:.2rem .35rem !important; font-size:.9rem !important; border-radius:10px !important; }

/* メタ行（ID） */
.enabled-line{ display:flex; align-items:center; gap:8px; margin-top:4px; }
.enabled-badge{
  font-size: 11px; padding: 2px 8px; border-radius:999px; border:1px solid rgba(128,140,160,.35);
  background: rgba(150,170,200,.12);
}
[data-theme="dark"] .enabled-badge{ border-color: rgba(120,140,180,.35); background: rgba(120,140,180,.12); }

/* Expander を軽くカード風 */
[data-testid="stExpander"]{ border-radius:12px; border:1px solid rgba(128,128,128,.25); }

/* 行操作ボタンを少し詰める */
.step-toolbar button { padding:0.2rem 0.4rem; font-size:0.9rem }
</style>
""", unsafe_allow_html=True)


# ============ ユーティリティ ============
def safe_rerun():
    try:
        st.rerun()
    except Exception:
        try:
            st.experimental_rerun()
        except Exception:
            pass

def json_pretty(obj: Any) -> str:
    try:
        return json.dumps(obj, ensure_ascii=False, indent=2)
    except Exception:
        return str(obj)

def force_parse_json(text: str) -> Optional[Any]:
    if text is None:
        return None
    if "```" in text:
        parts = text.split("```")
        for i in range(1, len(parts), 2):
            body = parts[i]
            try:
                return json.loads(body)
            except Exception:
                pass
    stack = []
    start = None
    for i, ch in enumerate(text):
        if ch in "{[":
            if not stack:
                start = i
            stack.append(ch)
        elif ch in "}]":
            if not stack:
                continue
            op = stack.pop()
            if (op == "{" and ch != "}") or (op == "[" and ch != "]"):
                stack = []
                start = None
                continue
            if not stack and start is not None:
                cand = text[start:i+1]
                try:
                    return json.loads(cand)
                except Exception:
                    start = None
    return None

def normalize_unified(parsed: Any) -> Dict[str, Any]:
    if isinstance(parsed, dict):
        if any(k in parsed for k in ("ok", "message", "kv_patch", "meta")):
            return {
                "ok": bool(parsed.get("ok", True)),
                "message": parsed.get("message") or {},
                "kv_patch": parsed.get("kv_patch") or {},
                "meta": parsed.get("meta") or {},
            }
        return {"ok": True, "message": parsed, "kv_patch": {}, "meta": {}}
    elif parsed is None:
        return {"ok": False, "message": {"error": "parse_failed"}, "kv_patch": {}, "meta": {}}
    else:
        return {"ok": True, "message": {"value": parsed}, "kv_patch": {}, "meta": {}}


# ============ システムプロンプト（利用者不可） ============
SCHEMA_ENFORCER = (
    "以下の制約を厳守:\n"
    "- 出力は有効なJSONのみ（前置きやコードブロック禁止）\n"
    "- スキーマ:\n"
    "{ \"ok\": true, \"message\": {...}, \"kv_patch\": {...}, \"meta\": {...} }\n"
)


# ============ Markdown 目次生成 ============
def compute_doc_index(md: str) -> Dict[str, Any]:
    if not isinstance(md, str):
        md = str(md or "")
    lines = md.splitlines()
    flat = []
    stack = []
    toc = {"title": "ROOT", "children": []}
    stack.append((0, toc))
    anchor_counts: Dict[str, int] = {}

    def slugify(title: str) -> str:
        t = re.sub(r"[^\w\-一-龠ぁ-んァ-ンー]+", "-", title.strip()).strip("-").lower()
        if not t:
            t = "section"
        cnt = anchor_counts.get(t, 0) + 1
        anchor_counts[t] = cnt
        return t if cnt == 1 else f"{t}-{cnt}"

    for ln in lines:
        m = re.match(r"^(#{1,6})\s+(.+?)\s*$", ln)
        if not m:
            continue
        level = len(m.group(1))
        title = m.group(2).strip()
        anchor = slugify(title)
        node = {"title": title, "level": level, "anchor": anchor, "children": []}
        flat.append({"level": level, "title": title, "anchor": anchor})
        while stack and stack[-1][0] >= level:
            stack.pop()
        parent_level, parent_node = stack[-1]
        parent_node["children"].append(node)
        stack.append((level, node))
    return {"flat": flat, "toc": toc}


# ============ マクロ展開ユーティリティ ============
def _to_int_or_none(s: str) -> Optional[int]:
    try:
        return int(s)
    except Exception:
        return None

def deep_get(obj: Any, path: str) -> Any:
    if not path:
        return obj
    cur = obj
    for part in path.split("."):
        if cur is None:
            return None
        if isinstance(cur, dict):
            if part in cur:
                cur = cur[part]
            else:
                i = _to_int_or_none(part)
                if i is not None and str(i) in cur:
                    cur = cur[str(i)]
                else:
                    return None
        elif isinstance(cur, list):
            i = _to_int_or_none(part)
            if i is None or not (0 <= i < len(cur)):
                return None
            cur = cur[i]
        else:
            return None
    return cur

def _get_prev_row_cells(ctx: Dict[str, Any]) -> List[Dict[str, Any]]:
    resp = (ctx or {}).get("resp") or {}
    if not isinstance(resp, dict):
        return []
    prev = resp.get("-1") or []
    return prev if isinstance(prev, list) else []

def summarize_prev_row(ctx: Dict[str, Any]) -> List[Dict[str, Any]]:
    cells = _get_prev_row_cells(ctx)
    out = []
    for cell in cells:
        out.append({
            "name": cell.get("name"),
            "agent_id": cell.get("agent_id"),
            "message": cell.get("message"),
            "raw": cell.get("raw"),
        })
    return out

def find_judge_results(ctx: Dict[str, Any]) -> Any:
    resp = (ctx or {}).get("resp") or {}
    if not isinstance(resp, dict):
        return None
    keys = []
    for k in resp.keys():
        try:
            keys.append(int(k))
        except Exception:
            pass
    keys = sorted([k for k in keys if k < 0], reverse=True)
    results = []
    for k in keys:
        cells = resp.get(str(k)) or []
        for cell in cells:
            name = (cell.get("name") or "").lower()
            if "judge" in name:
                v = cell.get("message") or cell.get("raw")
                if v is not None:
                    results.append(v)
        if results:
            break
    if not results:
        return None
    return results[0] if len(results) == 1 else results

def _get_prev_first_response(ctx: Dict[str, Any]) -> Any:
    cells = _get_prev_row_cells(ctx)
    if not cells:
        return None
    cell = cells[0]
    return cell.get("message") or cell.get("raw")

def _get_prev_first_prompt_text(row_idx: int) -> Optional[str]:
    try:
        if row_idx <= 0:
            return None
        prev_row = st.session_state.grid[row_idx - 1]
        if not prev_row:
            return None
        ag_prev_first = prev_row[0]
        for item in reversed(ag_prev_first.history):
            if item.get("role") == "user":
                return item.get("content")
        return None
    except Exception:
        return None

def normalize_context_path(p: str) -> str:
    if not isinstance(p, str):
        return ""
    s = re.sub(r"\s+", "", p)
    def repl(m):
        g1, g2, g3 = m.group(1), m.group(2), m.group(3)
        if g1 is not None: return "." + g1
        if g2 is not None: return "." + g2
        if g3 is not None: return "." + g3
        return ""
    s = re.sub(r'\[(?:"([^"]+)"|\'([^\']+)\'|(-?\d+))\]', repl, s)
    s = re.sub(r"\.+", ".", s).lstrip(".")
    return s

def expand_prompt_macros(prompt: str, ctx: Dict[str, Any], row_idx: Optional[int] = None) -> str:
    if prompt is None:
        return ""
    s = prompt.replace("｛", "{").replace("｝", "}")

    def _repl_prev_agent(m):
        agent_id = m.group(1)
        for cell in _get_prev_row_cells(ctx):
            if cell.get("agent_id") == agent_id:
                v = cell.get("message") or cell.get("raw") or {}
                return json_pretty(v) if isinstance(v, (dict, list)) else str(v)
        return f"(PrevStep.Agent:{agent_id} not found)"

    def _repl_prev_all(_m):
        return json_pretty(summarize_prev_row(ctx))

    def _repl_target_doc(_m):
        v = deep_get(ctx, "global_kv.target_doc")
        return v if isinstance(v, str) else (json_pretty(v) if v is not None else "")

    def _repl_judge(_m):
        v = find_judge_results(ctx)
        return json_pretty(v) if isinstance(v, (dict, list)) else (v or "(JudgeResult not found)")

    def _repl_prev_first_prompt(_m):
        txt = _get_prev_first_prompt_text(row_idx or 0)
        return txt or "(PrevStep.First.Prompt not found)"

    def _repl_prev_first_resp(_m):
        v = _get_prev_first_response(ctx)
        return json_pretty(v) if isinstance(v, (dict, list)) else (v or "(PrevStep.First.Response not found)")

    def _repl_context_path(m):
        path = normalize_context_path(m.group(1))
        v = deep_get(ctx, path)
        return json_pretty(v) if isinstance(v, (dict, list)) else (v if v is not None else f"(not found: Context.{path})")

    s = re.sub(r"\{PrevStep\.Agent:([A-Za-z0-9\-]+)\}", _repl_prev_agent, s)
    s = re.sub(r"\{PrevStep\.All\}", _repl_prev_all, s)
    s = re.sub(r"\{TargetDoc\}", _repl_target_doc, s)
    s = re.sub(r"\{JudgeResult\}", _repl_judge, s)
    s = re.sub(r"\{PrevStep\.First\.Prompt\}", _repl_prev_first_prompt, s)
    s = re.sub(r"\{PrevStep\.First\.Response\}", _repl_prev_first_resp, s)
    s = re.sub(r"\{Context\.([^\}]+)\}", _repl_context_path, s)
    return s


# ============ OpenAI ============
def get_openai_api_key() -> str:
    key = st.session_state.get("OPENAI_API_KEY") or os.environ.get("OPENAI_API_KEY") or os.environ.get("LLM_API_KEY")
    if not key:
        s_path = os.path.join(os.path.dirname(__file__), "secrets.toml")
        if os.path.exists(s_path):
            try:
                try:
                    import tomllib
                    with open(s_path, "rb") as fh:
                        data = tomllib.load(fh)
                except Exception:
                    try:
                        import toml as _toml
                        with open(s_path, "r", encoding="utf-8") as fh:
                            data = _toml.loads(fh.read())
                    except Exception:
                        data = {}
                key = key or data.get("LLM_API_KEY") or data.get("OPENAI_API_KEY")
            except Exception:
                key = key
    if key:
        try:
            st.session_state.OPENAI_API_KEY = key
        except Exception:
            pass
    if not key:
        raise RuntimeError("OPENAI_API_KEY が未設定です（サイドバーで設定または secrets.toml に LLM_API_KEY を設定してください）")
    return key

def call_llm(messages: List[Dict[str, str]], model: str, temperature: float, max_tokens: int, seed: Optional[int]) -> str:
    import openai
    key = get_openai_api_key()
    os.environ.setdefault("OPENAI_API_KEY", key)
    params = {"model": model, "temperature": float(temperature), "max_tokens": int(max_tokens)}
    if seed is not None:
        params["seed"] = int(seed)
    last_exc = None
    OpenAIClient = getattr(openai, "OpenAI", None)
    if OpenAIClient:
        try:
            try:
                client = OpenAIClient(api_key=key) if callable(OpenAIClient) else OpenAIClient()
            except Exception:
                client = OpenAIClient()
            try:
                chat = getattr(client, "chat", None)
                if chat and hasattr(chat, "completions"):
                    resp = chat.completions.create(messages=messages, **params)
                else:
                    responses = getattr(client, "responses", None)
                    if responses and hasattr(responses, "create"):
                        resp = responses.create(messages=messages, **params)
                    else:
                        resp = None
                if resp is not None:
                    try:
                        choice = resp.choices[0]
                        content = None
                        if hasattr(choice, "message"):
                            msg = choice.message
                            content = getattr(msg, "content", None)
                        elif isinstance(choice, dict):
                            msg = choice.get("message")
                            if isinstance(msg, dict):
                                content = msg.get("content")
                            else:
                                content = choice.get("text") or choice.get("content")
                        return str(content) if content is not None else str(resp)
                    except Exception as e:
                        last_exc = e
            except Exception as e:
                last_exc = e
        except Exception as e:
            last_exc = e
    try:
        openai.api_key = key
        resp = openai.ChatCompletion.create(model=model, messages=messages, temperature=float(temperature), max_tokens=int(max_tokens))
        try:
            return resp.choices[0].message["content"]
        except Exception:
            try:
                return resp.choices[0].text
            except Exception:
                return str(resp)
    except Exception as e:
        last_exc = e
    if last_exc:
        raise last_exc
    raise RuntimeError("openai client not available")


# ============ データ構造 ============
@dataclass
class Agent:
    id: str
    name: str = "Agent"
    user_prompt: str = ""
    model: str = "gpt-4o-mini"
    temperature: float = 0.3
    max_tokens: int = 1200
    seed: Optional[int] = None
    enabled: bool = True

    history: List[Dict[str, str]] = field(default_factory=list)
    last_raw: str = ""
    last_json: Any = None

def new_agent(name: str = "Agent") -> Agent:
    return Agent(id=str(uuid.uuid4())[:8], name=name)


# ============ アプリ状態初期化 ============
def ensure_state():
    if "OPENAI_API_KEY" not in st.session_state:
        st.session_state.OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY", "")
    if "grid" not in st.session_state:
        st.session_state.grid = [[new_agent("Ingestor"), new_agent("Editor"), new_agent("Completer"), new_agent("Optimizer")]]
    if "current_row" not in st.session_state:
        st.session_state.current_row = 0
    if "loop_index" not in st.session_state:
        st.session_state.loop_index = 0
    if "current_loop_steps" not in st.session_state:
        st.session_state.current_loop_steps = []
    if "history_loops" not in st.session_state:
        st.session_state.history_loops = []
    if "view" not in st.session_state:
        st.session_state.view = "main"
    if "global_kv" not in st.session_state:
        base_doc = (
            "# タイトル\n\n"
            "## 概要\n\n(ここに概要)\n\n"
            "## 前提条件\n\n- \n\n"
            "## 環境セットアップ\n\n- \n\n"
            "## 手順\n\n"
            "### ステップ1\n\n(やること)\n\n"
            "### ステップ2\n\n(やること)\n\n"
            "### ステップ3\n\n(やること)\n\n"
            "## よくあるエラーと対処\n\n"
            "- 症状: \n  - 原因: \n  - 対処: \n\n"
            "## ベストプラクティス\n\n- \n\n"
            "## 次のステップ\n\n- \n\n"
            "## 参考資料\n\n- \n\n"
            "## 付録\n\n- \n"
        )
        desired_outline_default = [
            "# タイトル",
            "概要",
            "前提条件",
            "環境セットアップ",
            "手順",
            "よくあるエラーと対処",
            "ベストプラクティス",
            "次のステップ",
            "参考資料",
            "付録",
        ]
        st.session_state.global_kv = {
            "doc_title": "ターゲット文書",
            "target_doc": base_doc,
            "doc_index": compute_doc_index(base_doc),
            "sources_word": [],
            "sources_excel": [],
            "sources_csv": [],
            "sources_text": [],
            "sources_web": [],
            "desired_outline": desired_outline_default,
        }
    if "defaults" not in st.session_state:
        st.session_state.defaults = {
            "model": "gpt-4o-mini",
            "temperature": 0.3,
            "max_tokens": 1200,
            "seed": None,
            "columns_per_row": 3,
        }
    if "ui_target_agent_id" not in st.session_state:
        st.session_state.ui_target_agent_id = None
    # 非表示にするエージェントID（指定があれば）
    if "hidden_agent_ids" not in st.session_state:
        st.session_state.hidden_agent_ids = ["26bd9a2b"]  # 指定があればここに列挙

ensure_state()


# ============ Context 構築 ============
def build_indexed_response_current(row_idx: int) -> Dict[str, List[Dict[str, Any]]]:
    mapping: Dict[str, List[Dict[str, Any]]] = {}
    for abs_row in range(0, row_idx):
        rel = abs_row - row_idx
        row = st.session_state.grid[abs_row]
        cells: List[Dict[str, Any]] = []
        for c, agent in enumerate(row):
            last_user = None
            try:
                for h in reversed(agent.history):
                    if h.get("role") == "user":
                        last_user = h.get("content")
                        break
            except Exception:
                last_user = None
            cells.append({
                "message": agent.last_json,
                "raw": agent.last_raw or None,
                "agent_id": agent.id,
                "name": agent.name,
                "row": abs_row,
                "col": c,
                "prompt": last_user,
            })
        mapping[str(rel)] = cells
    return mapping

def build_context_for_row(row_idx: int) -> Dict[str, Any]:
    resp_cur = build_indexed_response_current(row_idx)
    prompt_map: Dict[str, List[Optional[str]]] = {}
    resp_only_map: Dict[str, List[Optional[str]]] = {}
    for k, cells in (resp_cur or {}).items():
        if isinstance(cells, list):
            prompt_map[k] = [(cell.get("prompt") if isinstance(cell, dict) else None) for cell in cells]
            resp_only_map[k] = [(cell.get("raw") if isinstance(cell, dict) else None) for cell in cells]
        else:
            prompt_map[k] = []
            resp_only_map[k] = []
    return {
        "loop_index": st.session_state.loop_index,
        "current_row": row_idx,
        "response": {"0": resp_cur},
        "resp": resp_cur,
        "Prompt": prompt_map,
        "RespOnly": resp_only_map,
        "System": SCHEMA_ENFORCER,
        "global_kv": st.session_state.global_kv,
    }


# ============ Workflow 保存/読込 ============
def serialize_workflow() -> str:
    grid_dump: List[List[Dict[str, Any]]] = []
    for row in st.session_state.grid:
        row_dump = []
        for ag in row:
            row_dump.append({
                "id": ag.id,
                "name": ag.name,
                "user_prompt": ag.user_prompt,
                "model": ag.model,
                "temperature": ag.temperature,
                "max_tokens": ag.max_tokens,
                "seed": ag.seed,
                "enabled": ag.enabled,
            })
        row_dump.append
        grid_dump.append(row_dump)
    data = {
        "version": 2,
        "defaults": st.session_state.defaults,
        "global_kv": st.session_state.global_kv,
        "grid": grid_dump,
        "current_row": st.session_state.current_row,
        "loop_index": st.session_state.loop_index,
    }
    return json.dumps(data, ensure_ascii=False, indent=2)

def deserialize_workflow(text: str):
    try:
        data = json.loads(text)
    except Exception as e:
        st.error(f"読込失敗: {e}")
        return
    try:
        st.session_state.defaults = data.get("defaults", st.session_state.defaults)
        st.session_state.global_kv = data.get("global_kv", st.session_state.global_kv)
        grid_dump = data.get("grid", [])
        grid: List[List[Agent]] = []
        for row in grid_dump:
            row_agents: List[Agent] = []
            for a in row:
                ag = Agent(
                    id=a.get("id") or str(uuid.uuid4())[:8],
                    name=a.get("name", "Agent"),
                    user_prompt=a.get("user_prompt", ""),
                    model=a.get("model", st.session_state.defaults.get("model", "gpt-4o-mini")),
                    temperature=a.get("temperature", 0.3),
                    max_tokens=a.get("max_tokens", 1200),
                    seed=a.get("seed"),
                    enabled=a.get("enabled", True),
                )
                row_agents.append(ag)
            grid.append(row_agents)
        if grid:
            st.session_state.grid = grid
        st.session_state.current_row = int(data.get("current_row", 0))
        st.session_state.loop_index = int(data.get("loop_index", 0))
        st.success("ワークフローを読み込みました")
    except Exception as e:
        st.error(f"ワークフロー変換失敗: {e}")


# ============ Agent 実行 ============
def run_agent(agent: Agent, row_idx: int) -> Agent:
    if not agent.enabled:
        return agent

    ctx = build_context_for_row(row_idx)
    expanded_instr = expand_prompt_macros(agent.user_prompt or "", ctx, row_idx=row_idx)

    messages = [{"role": "system", "content": SCHEMA_ENFORCER}] + agent.history[-20:] + [{"role": "user", "content": expanded_instr}]

    model = agent.model or st.session_state.defaults["model"]
    temperature = agent.temperature if agent.temperature is not None else st.session_state.defaults["temperature"]
    max_tokens = int(agent.max_tokens or st.session_state.defaults["max_tokens"])
    seed = agent.seed if agent.seed is not None else st.session_state.defaults["seed"]

    try:
        raw = call_llm(messages, model=model, temperature=temperature, max_tokens=max_tokens, seed=seed)
    except Exception as e:
        raw = json.dumps({"ok": False, "message": {"error": f"LLM error: {e}"}, "kv_patch": {}, "meta": {"event": "error"}}, ensure_ascii=False)

    parsed = force_parse_json(raw)
    unified = normalize_unified(parsed)

    try:
        if "meta" not in unified or not isinstance(unified["meta"], dict):
            unified["meta"] = {}
        if "event" not in unified["meta"]:
            unified["meta"]["event"] = "success" if unified.get("ok") else "error"
    except Exception:
        pass

    agent.history.append({"role": "user", "content": expanded_instr})
    agent.history.append({"role": "assistant", "content": raw})
    agent.history = agent.history[-40:]

    agent.last_raw = raw
    agent.last_json = unified
    return agent


# ============ ワークフロー実行 ============
def _apply_prompt_patch(prompt_patch: Dict[str, Any], current_row: int):
    if not isinstance(prompt_patch, dict):
        return
    grid = st.session_state.grid
    for off_key, col_map in prompt_patch.items():
        try:
            off = int(off_key)
        except Exception:
            continue
        tgt_row = current_row + off
        if not (0 <= tgt_row < len(grid)):
            continue
        if not isinstance(col_map, dict):
            continue
        for c_key, new_prompt in col_map.items():
            try:
                cidx = int(c_key)
            except Exception:
                continue
            if not (0 <= cidx < len(grid[tgt_row])):
                continue
            if not isinstance(new_prompt, str):
                continue
            ag = grid[tgt_row][cidx]
            ag.user_prompt = new_prompt
            st.session_state.grid[tgt_row][cidx] = ag
            try:
                pk = f"det_prompt_{ag.id}_pending"
                st.session_state[pk] = new_prompt
                safe_rerun()
            except Exception:
                pass

def _apply_kv_patch(patch: Dict[str, Any]):
    if not isinstance(patch, dict):
        return
    for k, v in patch.items():
        if k == "Prompt":
            continue
        if k == "target_doc" and isinstance(v, str):
            st.session_state.global_kv["target_doc"] = v
            try:
                st.session_state.global_kv["doc_index"] = compute_doc_index(v)
            except Exception:
                pass
        else:
            st.session_state.global_kv[k] = v

def _parse_next_spec(next_spec: Any, current_row: int) -> Optional[Any]:
    if not next_spec:
        return None
    if isinstance(next_spec, dict):
        if next_spec.get("end"):
            return "end"
        if "row" in next_spec:
            r = next_spec["row"]
            if isinstance(r, str) and (r.startswith("+") or r.startswith("-")):
                try:
                    return current_row + int(r)
                except Exception:
                    return None
            if isinstance(r, int):
                return r
    if isinstance(next_spec, str):
        if next_spec.lower() in ("end", "finish"):
            return "end"
        if next_spec.startswith("+") or next_spec.startswith("-"):
            try:
                return current_row + int(next_spec)
            except Exception:
                return None
        try:
            return int(next_spec)
        except Exception:
            return None
    return None

def step_execute(row_idx: int):
    if row_idx >= len(st.session_state.grid):
        return
    row = st.session_state.grid[row_idx]
    step_result: Dict[str, Any] = {}

    for col, agent in enumerate(row):
        if not agent.enabled:
            continue
        updated = run_agent(agent, row_idx)
        st.session_state.grid[row_idx][col] = updated
        step_result[updated.id] = updated.last_json
        if isinstance(updated.last_json, dict):
            patch = updated.last_json.get("kv_patch") or {}
            _apply_kv_patch(patch)
            if isinstance(patch.get("Prompt"), dict):
                _apply_prompt_patch(patch.get("Prompt") or {}, row_idx)
            try:
                if isinstance(patch.get("target_doc"), str):
                    updated.last_json.setdefault("kv_patch", {})["target_doc"] = ""
                if isinstance(patch.get("Prompt"), dict):
                    updated.last_json.setdefault("kv_patch", {})["Prompt"] = {}
                st.session_state.grid[row_idx][col] = updated
            except Exception:
                pass

    if len(st.session_state.current_loop_steps) == row_idx:
        st.session_state.current_loop_steps.append(step_result)
    else:
        while len(st.session_state.current_loop_steps) < row_idx:
            st.session_state.current_loop_steps.append({})
        st.session_state.current_loop_steps.append(step_result)

    next_ptr = None
    try:
        for ag in st.session_state.grid[row_idx]:
            uni = ag.last_json if isinstance(ag.last_json, dict) else None
            if not uni:
                continue
            meta = uni.get("meta") or {}
            cand = _parse_next_spec(meta.get("next"), row_idx)
            if cand is not None:
                next_ptr = cand
                break
    except Exception:
        next_ptr = None

    if next_ptr == "end":
        st.session_state.history_loops.append({"loop_index": st.session_state.loop_index, "steps": st.session_state.current_loop_steps})
        st.session_state.loop_index += 1
        st.session_state.current_loop_steps = []
        st.session_state.current_row = 0
        return

    if isinstance(next_ptr, int) and 0 <= next_ptr < len(st.session_state.grid):
        st.session_state.current_row = next_ptr
        return

    if row_idx == len(st.session_state.grid) - 1:
        st.session_state.history_loops.append({"loop_index": st.session_state.loop_index, "steps": st.session_state.current_loop_steps})
        st.session_state.loop_index += 1
        st.session_state.current_loop_steps = []
        st.session_state.current_row = 0
    else:
        st.session_state.current_row = row_idx + 1

def run_one_step():
    if not st.session_state.grid or not st.session_state.grid[0]:
        st.warning("エージェントがありません")
        return
    step_execute(st.session_state.current_row)

def run_to_end():
    if not st.session_state.grid or not st.session_state.grid[0]:
        st.warning("エージェントがありません")
        return
    while True:
        row_idx = st.session_state.current_row
        step_execute(row_idx)
        if st.session_state.current_row == 0 and row_idx == len(st.session_state.grid) - 1:
            break

def run_to_end_times(n: int):
    n = max(1, int(n or 1))
    for _ in range(n):
        run_to_end()


# ============ 画面遷移 / 検索 ============
def go_detail(agent_id: str):
    st.session_state.ui_target_agent_id = agent_id
    st.session_state.view = "detail"
    safe_rerun()

def go_main():
    st.session_state.view = "main"
    st.session_state.ui_target_agent_id = None
    safe_rerun()

def find_agent_pos(agent_id: str):
    for r, row in enumerate(st.session_state.grid):
        for c, agent in enumerate(row):
            if agent.id == agent_id:
                return (r, c, agent)
    return None


# ============ 取込（Word/Excel/CSV/Text/Web） ============
def read_word_to_markdown(file) -> str:
    tmp_path = None
    out_md_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            data = file.read()
            tmp.write(data)
            tmp_path = tmp.name
        if pypandoc is not None:
            try:
                try:
                    pypandoc.get_pandoc_version()
                except OSError:
                    try:
                        pypandoc.download_pandoc()
                    except Exception:
                        pass
                out_md_path = tmp_path + '.md'
                pypandoc.convert_file(tmp_path, 'md', outputfile=out_md_path, extra_args=['--standalone'])
                with open(out_md_path, 'r', encoding='utf-8') as fh:
                    md = fh.read()
                md = re.sub(r" {2,}", " ", md)
                return md
            except Exception:
                pass
        if Document is not None:
            try:
                doc = Document(tmp_path)
                parts = []
                for p in doc.paragraphs:
                    txt = (p.text or "").strip()
                    if txt:
                        parts.append(txt)
                for tbl in doc.tables:
                    for row in tbl.rows:
                        cells = [" " + (c.text or "").replace("\n", " ").strip() + " " for c in row.cells]
                        parts.append("|" + "|".join(cells) + "|")
                md = "\n\n".join(parts)
                md = re.sub(r" {2,}", " ", md)
                return md
            except Exception:
                pass
        try:
            sample = data[:2000] if isinstance(data, (bytes, bytearray)) else str(data)[:2000]
            return f"```\n{sample}\n```\n(注: docx 解析に失敗)"
        except Exception:
            return "(docx の解析に失敗)"
    finally:
        try:
            if tmp_path and os.path.exists(tmp_path):
                os.remove(tmp_path)
        except Exception:
            pass
        try:
            if out_md_path and os.path.exists(out_md_path):
                os.remove(out_md_path)
        except Exception:
            pass

def df_to_markdown_or_csv(df) -> str:
    if df is None:
        return ""
    try:
        return df.to_markdown(index=False)
    except Exception:
        return df.to_csv(index=False)

def read_excel_to_markdown(file) -> str:
    if pd is None:
        return "(pandas が無いためExcelを処理できません)"
    tmp_path = None
    try:
        suffix = Path(getattr(file, 'name', '') or '').suffix.lower() or '.xlsx'
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(file.read())
            tmp_path = tmp.name
        engine = None
        if suffix in ('.xlsx', '.xlsm'):
            engine = 'openpyxl'
        elif suffix == '.xls':
            engine = 'xlrd'
        try:
            if engine:
                xls = pd.ExcelFile(tmp_path, engine=engine)
            else:
                xls = pd.ExcelFile(tmp_path)
        except Exception:
            xls = pd.ExcelFile(tmp_path)
        out_parts = []
        for sheet in xls.sheet_names:
            try:
                df = xls.parse(sheet, header=None).fillna("")
                try:
                    md = df.to_markdown(index=False)
                except Exception:
                    md = df.to_csv(index=False)
                md = re.sub(r" {2,}", " ", md)
                out_parts.append(f"## {sheet}\n\n{md}")
            except Exception:
                out_parts.append(f"## {sheet}\n\n(シートの解析に失敗しました)")
        return "\n\n".join(out_parts)
    except Exception as e:
        return f"(Excel 解析失敗: {e})"
    finally:
        try:
            if tmp_path and os.path.exists(tmp_path):
                os.remove(tmp_path)
        except Exception:
            pass

def read_csv_to_markdown(file) -> str:
    if pd is None:
        try:
            data = file.read().decode("utf-8", errors="ignore")
            return f"```\n{data}\n```"
        except Exception as e:
            return f"(CSV 読み込み失敗: {e})"
    try:
        df = pd.read_csv(file)
        return df_to_markdown_or_csv(df)
    except Exception as e:
        return f"(CSV 解析失敗: {e})"

def read_text(file) -> str:
    try:
        return file.read().decode("utf-8", errors="ignore")
    except Exception:
        try:
            return file.read().decode("cp932", errors="ignore")
        except Exception as e:
            return f"(テキスト読込失敗: {e})"

def html_to_markdown(html: str) -> str:
    if not isinstance(html, str):
        html = str(html or "")
    if _h2t:
        conv = _h2t.HTML2Text()
        conv.ignore_links = False
        conv.body_width = 0
        md = conv.handle(html)
    else:
        md = re.sub(r"<[^>]+>", "", html)
    md = re.sub(r"\n{3,}", "\n\n", md).strip()
    return md

def read_web_to_markdown(url: str) -> Dict[str, Any]:
    """
    指定URLから本文を抽出してMarkdown文字列を返す
    返却: { 'ok': bool, 'title': str, 'url': str, 'content_md': str, 'fetched_at': str, 'error': str|None }
    """
    from datetime import datetime, timezone

    if not url or not isinstance(url, str):
        return {"ok": False, "title": "", "url": url, "content_md": "", "fetched_at": "", "error": "invalid_url"}

    title = ""
    md = ""
    err = None

    try:
        # 1) trafilatura を優先
        if trafilatura is not None:
            downloaded = trafilatura.fetch_url(url)
            if downloaded:
                extracted = trafilatura.extract(
                    downloaded,
                    include_formatting=True,
                    include_links=True,
                    favor_recall=True,
                )
                if extracted:
                    md = extracted
                    lines = [ln.strip() for ln in md.splitlines() if ln.strip()]
                    if lines:
                        title = lines[0].lstrip("# ").strip()[:80]

        # 2) readability-lxml
        if not md and (requests is not None) and (ReadabilityDocument is not None):
            headers = {"User-Agent": "Mozilla/5.0 (compatible; ContentIngestor/1.0; +https://example.com)"}
            r = requests.get(url, headers=headers, timeout=20)
            r.raise_for_status()
            html = r.text
            doc = ReadabilityDocument(html)
            title = (doc.short_title() or "").strip()[:80] or title
            cleaned_html = doc.summary(html_partial=True)
            md = html_to_markdown(cleaned_html)

        # 3) BeautifulSoup フォールバック
        if not md and (requests is not None) and (BeautifulSoup is not None):
            headers = {"User-Agent": "Mozilla/5.0 (compatible; ContentIngestor/1.0; +https://example.com)"}
            r = requests.get(url, headers=headers, timeout=20)
            r.raise_for_status()
            html = r.text
            soup = BeautifulSoup(html, "html.parser")
            # ノイズ除去
            for tag in soup(["script", "style", "noscript", "svg", "canvas", "form", "iframe"]):
                tag.decompose()
            for tag in soup.find_all(["header", "footer", "nav", "aside"]):
                tag.decompose()
            # タイトル
            if soup.title and soup.title.string:
                title = (soup.title.string or "").strip()[:80] or title
            # 本文候補
            main = soup.find(["main", "article"]) or soup.body or soup
            cleaned_html = str(main)
            md = html_to_markdown(cleaned_html)

        # 軽いノイズ抑制
        if md:
            noise_patterns = [
                r"^この記事をシェア", r"^関連記事", r"^広告", r"^スポンサー", r"^同意します", r"^Cookie",
            ]
            lines = md.splitlines()
            cleaned = []
            for ln in lines:
                if any(re.search(pat, ln, re.I) for pat in noise_patterns):
                    continue
                cleaned.append(ln)
            md = "\n".join(cleaned).strip()

        if not md:
            err = "extract_failed"

    except Exception as e:
        err = f"{e}"

    return {
        "ok": bool(md),
        "title": title or "(no title)",
        "url": url,
        "content_md": md or "",
        "fetched_at": datetime.now(timezone.utc).isoformat(),
        "error": err,
    }


# ============ UI: Sidebar（保存/読込/設定/取込/実行） ============
def apply_default_knowledge_flow():
    st.session_state.grid = [[
        Agent(id=str(uuid.uuid4())[:8], name="Ingestor", user_prompt=(
            "あなたは情報インジェスト担当です。\n"
            "目的:\n"
            "- 取り込みソース（Word/Excel/CSV/Text/Web）を正規化し、既存の本文に統合します。\n"
            "- 本文を「チュートリアル向けMarkdownアウトライン」に沿って構成整理します。\n"
            "入力:\n"
            "- 本文（現状）: {Context.global_kv.target_doc}\n"
            "- 取り込み（Word）: {Context.global_kv.sources_word}\n"
            "- 取り込み（Excel）: {Context.global_kv.sources_excel}\n"
            "- 取り込み（CSV）: {Context.global_kv.sources_csv}\n"
            "- 取り込み（Text）: {Context.global_kv.sources_text}\n"
            "- 取り込み（Web）: {Context.global_kv.sources_web}\n"
            "- 望ましいアウトライン（任意）: {Context.global_kv.desired_outline}\n"
            "やること:\n"
            "- ソースを読み、重複・ノイズを除去して要点をMarkdownに正規化。\n"
            "- desired_outline が空なら、以下の既定アウトラインで本文を再構成:\n"
            "  「タイトル/概要/前提条件/環境セットアップ/手順(ステップ1..)/よくあるエラーと対処/ベストプラクティス/次のステップ/参考資料/付録」\n"
            "- 既存本文の有用部は保持し、必要に応じて節（## 見出し）を新設して統合。\n"
            "- 根拠は取り込みソースからのみ。不明点は推測せず「TODO: ～を確認」と明示。\n"
            "- 参考リンクはsources_webからのみ引用（存在する場合）。出典は「参考資料」節に箇条書きで集約。\n"
            "- 文体は簡潔・手順中心。見出し・箇条書き・コードブロック（必要時）の整形。\n"
            "- タイトル未設定なら適切に補う。\n"
            "- token制限内で最重要項目を優先統合。\n"
            "失敗時の振る舞い:\n"
            "- 新規の有益情報が無い場合は本文を大きく変更せず、notesに理由を記述。\n"
            "出力はJSONのみ（前置きやコードブロック禁止）:\n"
            "{\n"
            "  \"ok\": true,\n"
            "  \"message\": {\n"
            "    \"type\": \"ingest\",\n"
            "    \"notes\": \"(統合メモ: 何を取り込み、何を据え置いたか。TODOも明記)\",\n"
            "    \"changes\": [\"(追加/修正した主な見出しや節)\"]\n"
            "  },\n"
            "  \"kv_patch\": {\"target_doc\": \"(統合後の本文Markdown全文: チュートリアルアウトラインに準拠)\"},\n"
            "  \"meta\": {}\n"
            "}\n"
        )),
        Agent(id=str(uuid.uuid4())[:8], name="Editor", user_prompt=(
            "あなたは本文編集担当です。\n"
            "目的:\n"
            "- 本文を読みやすく再編集し、チュートリアル向けMarkdownアウトラインに完全準拠させます。\n"
            "入力:\n"
            "- 現在の本文: {Context.global_kv.target_doc}\n"
            "- 本文インデックス: {Context.global_kv.doc_index}\n"
            "- 直前STEPの出力: {Context.resp[\"-1\"]}\n"
            "- 望ましいアウトライン（任意）: {Context.global_kv.desired_outline}\n"
            "やること:\n"
            "- desired_outline があればそれに厳密に合わせ、無ければ既定アウトラインに準拠。\n"
            "- 節構成の重複を解消、見出し粒度を統一、ステップ名を動詞から始まる形に整理。\n"
            "- 用語を統一、冗長表現や曖昧表現を削減。箇条書き・表・コードブロックを適切に使用。\n"
            "- 「よくあるエラーと対処」は 症状→原因→対処 の順で最小テンプレに整形。\n"
            "- 出典やURLは「参考資料」節に集約し、本文中は最小限の参照。\n"
            "- 重要情報は残し、重複のみ統合・省略。\n"
            "出力はJSONのみ（前置きやコードブロック禁止）:\n"
            "{\n"
            "  \"ok\": true,\n"
            "  \"message\": {\"type\": \"commit\", \"summary\": \"(変更点の要約: 構成変更、表現統一、重複解消など)\"},\n"
            "  \"kv_patch\": {\"target_doc\": \"(編集後の本文全体: アウトライン準拠・読みやすさ改善済み)\"},\n"
            "  \"meta\": {}\n"
            "}\n"
        )),
        Agent(id=str(uuid.uuid4())[:8], name="Completer", user_prompt=(
            "あなたは構成補完担当です。\n"
            "目的:\n"
            "- 抜けている節や記述を検出し、必要最小限の内容を追加して完全性を高めます。\n"
            "入力:\n"
            "- 本文インデックス: {Context.global_kv.doc_index}\n"
            "- 本文: {Context.global_kv.target_doc}\n"
            "- 望ましいアウトライン（任意）: {Context.global_kv.desired_outline}\n"
            "やること:\n"
            "- desired_outline または既定アウトラインと照合し、欠落節を特定。\n"
            "- 「前提条件」「環境セットアップ」「手順」に必須の抜けがあれば優先補完。\n"
            "- 補完は最小限かつ実用重視。根拠不明の事項は「TODO: 要確認」を明記。\n"
            "- 既存記述と重複しないように差分追加。大幅書き換えは避ける。\n"
            "出力はJSONのみ（前置きやコードブロック禁止）:\n"
            "{\n"
            "  \"ok\": true,\n"
            "  \"message\": {\"type\": \"complete\", \"added\": [\"(追加した見出しや小節)\"], \"notes\": \"(補完方針や未解決TODO)\"},\n"
            "  \"kv_patch\": {\"target_doc\": \"(補完後の本文全体)\"},\n"
            "  \"meta\": {}\n"
            "}\n"
        )),
        Agent(id=str(uuid.uuid4())[:8], name="Optimizer", user_prompt=(
            "あなたはプロンプト改善担当です。\n"
            "目的:\n"
            "- 現在行の「左から1番＝Ingestor」のプロンプトを、取り込み漏れ防止・アウトライン徹底・非推測方針の明確化の観点で改善して即時更新します。\n"
            "入力:\n"
            "- 直前STEPの左から1番の出力要約: {Context.resp[\"-1\"][0].message}\n"
            "- グローバル情報（取り込みソース・アウトライン）: {Context.global_kv}\n"
            "チェック項目:\n"
            "- sources_word / sources_excel / sources_csv / sources_text / sources_web をすべて参照しているか\n"
            "- desired_outline を考慮（無い場合は既定アウトラインを明示）\n"
            "- 本文をチュートリアル向けMarkdownで構成\n"
            "- 未知はTODO明示・出典集約・token配慮・JSONのみ出力（コードブロック禁止）\n"
            "出力はJSONのみ（前置きやコードブロック禁止）:\n"
            "{\n"
            "  \"ok\": true,\n"
            "  \"message\": {\"type\": \"prompt_update\", \"before\": \"(現プロンプトの要約や改善観点)\", \"after\": \"(改善後プロンプト全文)\"},\n"
            "  \"kv_patch\": {\"Prompt\": {\"0\": {\"0\": \"(改善後プロンプト全文)\"}}},\n"
            "  \"meta\": {}\n"
            "}\n"
        )),
    ]]
    st.session_state.current_row = 0
    st.session_state.current_loop_steps = []
    st.success("デフォルトのナレッジフローを適用しました")

with st.sidebar:
    st.header("設定")
    st.session_state.OPENAI_API_KEY = st.text_input("OPENAI_API_KEY", value=st.session_state.get("OPENAI_API_KEY", ""), type="password")
    d = st.session_state.defaults
    d["model"] = st.text_input("Model", value=d.get("model", "gpt-4o-mini"))
    d["temperature"] = float(st.slider("temperature", 0.0, 2.0, float(d.get("temperature", 0.3)), 0.1))
    d["max_tokens"] = int(st.number_input("max_tokens", min_value=128, max_value=8192, value=int(d.get("max_tokens", 1200)), step=64))
    d["columns_per_row"] = int(st.number_input("表示列数", 1, 8, int(d.get("columns_per_row", 3)), 1))

    st.markdown("---")
    st.caption(f"Loop: {st.session_state.loop_index}")
    st.caption(f"Current Row: {st.session_state.current_row+1} / {len(st.session_state.grid)}")
    st.caption(f"Agents: {sum(len(r) for r in st.session_state.grid)}")

    st.markdown("---")
    if st.button("📐 デフォルトフロー適用"):
        apply_default_knowledge_flow()
        safe_rerun()

    save_json = serialize_workflow()
    st.download_button("💾 保存", data=save_json, file_name="workflow.json", mime="application/json")
    up_flow = st.file_uploader("📂 読込 (JSON)", type=["json"], key="wf_upload")
    if up_flow is not None:
        try:
            text = up_flow.read().decode("utf-8")
            if st.button("読込を適用"):
                deserialize_workflow(text)
                safe_rerun()
        except Exception as e:
            st.error(f"読込失敗: {e}")

    st.markdown("---")
    # ファイル取込（サイドバーに移動）
    st.subheader("ファイル取込")
    with st.expander("Word (.docx)"):
        wfile = st.file_uploader("Word ファイルを選択", type=["docx"], key="up_word")
        if wfile is not None:
            md = read_word_to_markdown(wfile)
            st.text_area("プレビュー（Markdown相当）", value=md, height=160, key="preview_word")
            if st.button("sources_word に保存", key="save_word"):
                st.session_state.global_kv["sources_word"] = [{"title": wfile.name, "content_md": md}]
                st.success("sources_word を更新しました")
    with st.expander("Excel (.xlsx/.xls/.xlsm)"):
        efile = st.file_uploader("Excel ファイルを選択", type=["xlsx", "xls", "xlsm"], key="up_excel")
        if efile is not None:
            md = read_excel_to_markdown(efile)
            st.text_area("プレビュー（Markdown/CSV）", value=md, height=160, key="preview_excel")
            if st.button("sources_excel に保存", key="save_excel"):
                st.session_state.global_kv["sources_excel"] = [{"title": efile.name, "content_md": md}]
                st.success("sources_excel を更新しました")
    with st.expander("CSV (.csv)"):
        cfile = st.file_uploader("CSV ファイルを選択", type=["csv"], key="up_csv")
        if cfile is not None:
            md = read_csv_to_markdown(cfile)
            st.text_area("プレビュー（Markdown/CSV）", value=md, height=160, key="preview_csv")
            if st.button("sources_csv に保存", key="save_csv"):
                st.session_state.global_kv["sources_csv"] = [{"title": getattr(cfile, "name", "data.csv"), "content_md": md}]
                st.success("sources_csv を更新しました")
    with st.expander("Text (.txt)"):
        tfile = st.file_uploader("Text ファイルを選択", type=["txt"], key="up_text")
        if tfile is not None:
            tx = read_text(tfile)
            st.text_area("プレビュー（テキスト）", value=tx, height=160, key="preview_text")
            if st.button("sources_text に保存", key="save_text"):
                st.session_state.global_kv["sources_text"] = [{"title": getattr(tfile, "name", "text.txt"), "content": tx}]
                st.success("sources_text を更新しました")
    with st.expander("Web (URL)"):
        url_text = st.text_area("URL（複数可・改行区切り）", value="", height=100, key="up_web_urls")
        if st.button("取得", key="fetch_web_btn"):
            urls = [u.strip() for u in (url_text or "").splitlines() if u.strip()]
            results = []
            for u in urls:
                res = read_web_to_markdown(u)
                results.append(res)
            st.session_state["web_fetch_results"] = results
            if results:
                preview = "\n\n---\n\n".join(
                    f"# {r.get('title')}\n\nURL: {r.get('url')}\n\n{r.get('content_md')[:2000]}{'... (truncated)' if len(r.get('content_md',''))>2000 else ''}"
                    for r in results
                )
                st.text_area("プレビュー（Markdown）", value=preview, height=200, key="preview_web")
        if st.button("sources_web に保存", key="save_web_btn"):
            results = st.session_state.get("web_fetch_results") or []
            items = [
                {"title": r.get("title"), "url": r.get("url"), "content_md": r.get("content_md"), "fetched_at": r.get("fetched_at")}
                for r in results if r.get("ok") and r.get("content_md")
            ]
            st.session_state.global_kv["sources_web"] = items
            st.success(f"sources_web を更新しました（{len(items)}件）")

    st.markdown("---")
    # 実行操作（サイドバーに移動）
    st.subheader("実行")
    if st.button("▶ Step 実行（現在行）"):
        run_one_step()
        st.success("Step 実行完了")
    times = st.number_input("最後まで実行 回数", min_value=1, max_value=100, value=1, step=1)
    if st.button("⏩ 最後まで実行（指定回数）"):
        run_to_end_times(times)
        st.success(f"{int(times)} 回実行完了")

    st.markdown("---")
    if st.button("履歴のみ初期化（Grid保持）"):
        for row in st.session_state.grid:
            for ag in row:
                ag.history.clear()
                ag.last_raw = ""
                ag.last_json = None
        st.session_state.current_row = 0
        st.session_state.current_loop_steps = []
        st.session_state.history_loops = []
        st.success("履歴初期化しました")
        safe_rerun()

    if st.button("全リセット（Grid含む）"):
        st.session_state.clear()
        ensure_state()
        st.success("初期化しました")
        safe_rerun()


# ============ ヘルパー（表示抑制: Completerなど） ============
def is_agent_hidden(ag: Agent) -> bool:
    if ag is None:
        return False
    # 指定IDで非表示
    if ag.id in set(st.session_state.get("hidden_agent_ids", [])):
        return True
    nm = (ag.name or "").strip().lower()
    # Completer/Complete 系は表示を隠す（枠も不要）
    if nm.startswith("completer") or nm.startswith("complete"):
        return True
    return False


# ============ UI: Main ============
def render_main():
    st.title("ワークフロー")

    base_cols = int(st.session_state.defaults["columns_per_row"])
    for r, row in enumerate(st.session_state.grid):
        st.subheader(f"Step {r+1}")
        tb1, tb2, _tb3 = st.columns([1,1,6], gap="small")
        with tb1:
            if st.button("⬆️ 行上", key=f"add_row_up_{r}"):
                st.session_state.grid.insert(r, [new_agent("Agent 1")])
                safe_rerun()
        with tb2:
            if st.button("⬇️ 行下", key=f"add_row_down_{r}"):
                st.session_state.grid.insert(r+1, [new_agent("Agent 1")])
                safe_rerun()

        # 可視エージェントのみ表示（Completer等は非表示）
        visible_row = [ag for ag in row if not is_agent_hidden(ag)]
        grid_cols = max(base_cols, len(visible_row))
        cols = st.columns(grid_cols)

        for c, ag in enumerate(visible_row):
            with cols[c]:
                is_current = (r == st.session_state.current_row)

                # カード（非表示対象はここまでで除外済み）
                card = st.container(border=True)
                with card:
                    # ヘッダー
                    bar_cls = "agent-topbar is-current" if is_current else "agent-topbar"
                    st.markdown(
                        f'<div class="{bar_cls}">'
                        f'  <div class="agent-title"><span class="agent-chip">S{r+1}-C{c+1}</span>{ag.name}</div>'
                        f'  <div style="display:flex; align-items:center; gap:8px;">'
                        f'    {"<span class=\'agent-badge-current\'>Current</span>" if is_current else ""}'
                        f'    <span style="font-size:12px; opacity:.8;">{ag.model}</span>'
                        f'  </div>'
                        f'</div>',
                        unsafe_allow_html=True
                    )

                    # 有効切替
                    en = st.checkbox("有効", value=ag.enabled, key=f"en_{ag.id}")
                    if en != ag.enabled:
                        ag.enabled = en
                        # grid 内のオリジナル参照を更新（行内の該当IDを探す）
                        for idx, orig in enumerate(st.session_state.grid[r]):
                            if orig.id == ag.id:
                                st.session_state.grid[r][idx] = ag
                                break

                    # メタ（Agent ID） はメイン画面から非表示

                    # ボタン行（主要ボタンのみ）
                    st.markdown('<div class="btn-row small-btn">', unsafe_allow_html=True)
                    bcols = st.columns(3, gap="small")
                    with bcols[0]:
                        if st.button("▶", key=f"run_{ag.id}", help="実行"):
                            updated = run_agent(ag, r)
                            if isinstance(updated.last_json, dict):
                                patch = updated.last_json.get("kv_patch") or {}
                                _apply_kv_patch(patch)
                                if isinstance(patch.get("Prompt"), dict):
                                    _apply_prompt_patch(patch.get("Prompt") or {}, r)
                                try:
                                    if isinstance(patch.get("target_doc"), str):
                                        updated.last_json.setdefault("kv_patch", {})["target_doc"] = ""
                                    if isinstance(patch.get("Prompt"), dict):
                                        updated.last_json.setdefault("kv_patch", {})["Prompt"] = {}
                                except Exception:
                                    pass
                            # grid 内のオリジナル参照を更新
                            for idx, orig in enumerate(st.session_state.grid[r]):
                                if orig.id == ag.id:
                                    st.session_state.grid[r][idx] = updated
                                    break
                    # '＋' ボタンはメイン画面から削除
                    with bcols[1]:
                        if st.button("－", key=f"del_{ag.id}", help="削除"):
                            # グリッド上の実位置を探して削除
                            for idx, orig in enumerate(st.session_state.grid[r]):
                                if orig.id == ag.id:
                                    st.session_state.grid[r].pop(idx)
                                    break
                            if not st.session_state.grid[r]:
                                st.session_state.grid.pop(r)
                            safe_rerun()
                    with bcols[2]:
                        if st.button("🔎", key=f"detail_{ag.id}", help="詳細"):
                            go_detail(ag.id)
                    st.markdown('</div>', unsafe_allow_html=True)

                    # 出力（Completerなどは非表示にしているためここは他エージェントのみ）
                    with st.expander("Parsed", expanded=False):
                        if ag.last_json:
                            st.code(json_pretty(ag.last_json.get("message")), language="json")
                        else:
                            st.info("(no parsed JSON)")
                    with st.expander("Raw", expanded=False):
                        if ag.last_raw:
                            st.code(ag.last_raw)
                        else:
                            st.info("(no raw)")

        if st.button("＋追加", key=f"add_tail_{r}"):
            st.session_state.grid[r].append(new_agent(f"Agent {len(st.session_state.grid[r])+1}"))
            safe_rerun()

    st.markdown("---")

    # ターゲット文書（Markdown）と目次をエージェントの下側に移動
    st.subheader("ターゲット文書（Markdown）")
    td = st.session_state.global_kv.get("target_doc", "")
    td_new = st.text_area("target_doc", value=td, height=200)
    if st.button("target_doc を保存"):
        st.session_state.global_kv["target_doc"] = td_new
        st.session_state.global_kv["doc_index"] = compute_doc_index(td_new)
        st.success("更新しました（目次再生成済）")
    with st.expander("目次（自動生成）", expanded=False):
        st.code(json_pretty(st.session_state.global_kv.get("doc_index")), language="json")

    st.markdown("---")
    st.subheader("現在ループの進捗")
    st.code(json_pretty(st.session_state.current_loop_steps), language="json")

    st.subheader("過去ループ履歴")
    depth = st.slider("表示件数", 0, 10, 3)
    view = st.session_state.history_loops[-depth:] if depth > 0 else []
    st.code(json_pretty(view), language="json")


# ============ UI: Detail ============
def render_detail():
    st.title("エージェント詳細")
    agent_id = st.session_state.get("ui_target_agent_id")
    if not agent_id:
        st.warning("エージェント未選択")
        if st.button("← メインへ"):
            go_main()
        return
    pos = find_agent_pos(agent_id)
    if not pos:
        st.error("エージェントが見つかりません")
        if st.button("← メインへ"):
            go_main()
        return
    r, c, ag = pos

    top1, top2 = st.columns([1,3])
    with top1:
        if st.button("← メインへ"):
            go_main()
    with top2:
        st.write(f"Step {r+1} / 列 {c+1}")

    ag.enabled = st.checkbox("有効", value=ag.enabled, key=f"det_en_{ag.id}")
    ag.name = st.text_input("名前", value=ag.name, key=f"det_name_{ag.id}")

    k = f"det_prompt_{ag.id}"
    pending_key = f"{k}_pending"
    # If a pending update exists (from a button click), apply it before creating the widget
    if pending_key in st.session_state:
        try:
            pending_val = st.session_state.pop(pending_key)
            # If the widget key already exists, avoid direct assignment (Streamlit forbids it).
            if k in st.session_state:
                # update agent object and grid so next render picks it up
                ag.user_prompt = pending_val
                st.session_state.grid[r][c] = ag
                safe_rerun()
            else:
                st.session_state[k] = pending_val
        except Exception:
            pass
    # Ensure the session_state prompt reflects the agent's prompt before widget creation
    if st.session_state.get(k, None) != (ag.user_prompt or ""):
        st.session_state[k] = ag.user_prompt or ""

    # 指示プロンプト（ユーザー）を大きくして戻るボタンの下側に配置
    ag.user_prompt = st.text_area("指示プロンプト（ユーザー）", key=k, height=480)
    ag.user_prompt = st.session_state.get(k, "")

    with st.expander("モデル設定", expanded=False):
        ag.model = st.text_input("model", value=ag.model, key=f"det_model_{ag.id}")
        ag.temperature = float(st.slider("temperature", 0.0, 2.0, float(ag.temperature), 0.1, key=f"det_temp_{ag.id}"))
        ag.max_tokens = int(st.number_input("max_tokens", 128, 8192, int(ag.max_tokens), 64, key=f"det_maxtok_{ag.id}"))
        seed_flag = st.checkbox("seedを使う", value=(ag.seed is not None), key=f"det_seed_f_{ag.id}")
        if seed_flag:
            ag.seed = int(st.number_input("seed", 0, 2**31-1, int(ag.seed or 0), 1, key=f"det_seed_v_{ag.id}"))
        else:
            ag.seed = None

    t_worker = (
        "あなたはプロンプト改善担当です。\n"
        "目的:\n"
        "- 現在行の「左から1番＝Ingestor」のプロンプトを、取り込み漏れ防止・アウトライン徹底・非推測方針の明確化の観点で改善して即時更新します。\n"
        "入力:\n"
        "- 直前STEPの左から1番の出力要約: {Context.resp[\"-1\"][0].message}\n"
        "- グローバル情報（取り込みソース・アウトライン）: {Context.global_kv}\n\n"
        "チェック項目:\n"
        "- sources_word / sources_excel / sources_csv / sources_text / sources_web をすべて参照しているか\n"
        "- desired_outline を考慮（無い場合は既定アウトラインを明示）\n"
        "- 本文をチュートリアル向けMarkdownで構成\n"
        "- 未知はTODO明示・出典集約・token配慮・JSONのみ出力（コードブロック禁止）\n\n"
        "出力はJSONのみ:\n"
        "{\n"
        "  \"ok\": true,\n"
        "  \"message\": {\"type\": \"prompt_update\", \"before\": \"(現プロンプトの要約や改善観点)\", \"after\": \"(改善後プロンプト全文)\"},\n"
        "  \"kv_patch\": {\"Prompt\": {\"0\": {\"0\": \"(改善後プロンプト全文)\"}}},\n"
        "  \"meta\": {}\n"
        "}\n"
    )

    t_editor = (
        "あなたは本文編集担当です。\n"
        "- 現在の本文: {Context.global_kv.target_doc}\n"
        "- 本文インデックス: {Context.global_kv.doc_index}\n"
        "- 直前STEPの出力: {Context.resp[\"-1\"]}\n"
        "- 望ましいアウトライン（任意）: {Context.global_kv.desired_outline}\n\n"
        "不足の節を補い、構成を整え、本文の完全版を返してください。JSONのみ:\n"
        "{\n"
        "  \"ok\": true,\n"
        "  \"message\": {\"type\": \"commit\", \"summary\": \"(変更点の要約)\"},\n"
        "  \"kv_patch\": {\"target_doc\": \"(修正後の本文全体)\"},\n"
        "  \"meta\": {}\n"
        "}\n"
    )

    bt1, bt2 = st.columns(2)
    with bt1:
        if st.button("テンプレ: Worker", key=f"tmplW_{ag.id}"):
            st.session_state[pending_key] = t_worker
            safe_rerun()
    with bt2:
        if st.button("テンプレ: Editor", key=f"tmplE_{ag.id}"):
            st.session_state[pending_key] = t_editor
            safe_rerun()

    st.markdown("#### 差し込みショートカット")

    def insert_at_cursor(snippet: str):
        base = st.session_state.get(k, "")
        if "{|}" in base:
            new_val = base.replace("{|}", snippet, 1)
        else:
            new_val = base + ("" if base.endswith("\n") or base == "" else "\n") + snippet
        # Store pending update and rerun so the pending value is applied before the text_area widget is created.
        st.session_state[pending_key] = new_val
        safe_rerun()

    rowA = st.columns(6)
    with rowA[0]:
        if st.button("📄 TDOC", key=f"ins_tdoc_{ag.id}", use_container_width=True):
            insert_at_cursor("{Context.global_kv.target_doc}")
    with rowA[1]:
        if st.button("📘 WORD", key=f"ins_word_{ag.id}", use_container_width=True):
            insert_at_cursor("{Context.global_kv.sources_word}")
    with rowA[2]:
        if st.button("📊 XLSX", key=f"ins_xlsx_{ag.id}", use_container_width=True):
            insert_at_cursor("{Context.global_kv.sources_excel}")
    with rowA[3]:
        if st.button("🧾 CSV ", key=f"ins_csv_{ag.id}", use_container_width=True):
            insert_at_cursor("{Context.global_kv.sources_csv}")
    with rowA[4]:
        if st.button("📜 TEXT", key=f"ins_text_{ag.id}", use_container_width=True):
            insert_at_cursor("{Context.global_kv.sources_text}")
    with rowA[5]:
        if st.button("🌐 WEB", key=f"ins_web_{ag.id}", use_container_width=True):
            insert_at_cursor("{Context.global_kv.sources_web}")

    rowB = st.columns(4)
    with rowB[0]:
        if st.button("📚 IDX ", key=f"ins_idx_{ag.id}", use_container_width=True):
            insert_at_cursor("{Context.global_kv.doc_index}")
    with rowB[1]:
        if st.button("👥 R-1 ", key=f"ins_r_1_{ag.id}", use_container_width=True):
            insert_at_cursor('{Context.resp["-1"]}')
    with rowB[2]:
        if st.button("💬 R1M ", key=f"ins_r1m_{ag.id}", use_container_width=True):
            insert_at_cursor('{Context.resp["-1"][0].message}')
    with rowB[3]:
        if st.button("🧠 PRM1", key=f"ins_kvpatch_prompt_prev1_{ag.id}", use_container_width=True):
            insert_at_cursor(
                '次のJSONで kv_patch.Prompt を必ず設定し、現在行の左から1番のプロンプト（Ingestor）を更新してください:\\n'
                '{\\n'
                '  "ok": true,\\n'
                '  "message": {"type": "prompt_update", "before": "(要約)", "after": "(改善後プロンプト)"},\\n'
                '  "kv_patch": {"Prompt": {"0": {"0": "(改善後プロンプト)"}}},\\n'
                '  "meta": {}\\n'
                '}'
            )

    ac1, ac2 = st.columns(2)
    with ac1:
        if st.button("▶ 実行", key=f"det_run_{ag.id}"):
            updated = run_agent(ag, r)
            if isinstance(updated.last_json, dict):
                patch = updated.last_json.get("kv_patch") or {}
                _apply_kv_patch(patch)
                if isinstance(patch.get("Prompt"), dict):
                    _apply_prompt_patch(patch.get("Prompt") or {}, r)
                try:
                    if isinstance(patch.get("target_doc"), str):
                        updated.last_json.setdefault("kv_patch", {})["target_doc"] = ""
                    if isinstance(patch.get("Prompt"), dict):
                        updated.last_json.setdefault("kv_patch", {})["Prompt"] = {}
                except Exception:
                    pass
            st.session_state.grid[r][c] = updated
            st.success("実行しました")
    with ac2:
        if st.button("⟳ クリア", key=f"det_clear_{ag.id}"):
            ag.last_raw = ""
            ag.last_json = None
            st.session_state.grid[r][c] = ag

    st.caption("Parsed (unified)")
    st.code(json_pretty(ag.last_json) if ag.last_json else "(なし)", language="json")
    st.caption("Raw")
    st.code(ag.last_raw or "(なし)")

    with st.expander("送信メッセージ確認", expanded=False):
        ctx = build_context_for_row(r)
        raw_instr = st.session_state.get(k, "")
        expanded_instr = expand_prompt_macros(raw_instr, ctx, row_idx=r)
        system_view = SCHEMA_ENFORCER
        user_view_before = raw_instr
        user_view_after = expanded_instr

        st.caption("System（システムプロンプト：編集不可）")
        st.code(system_view)
        st.caption("User（指示プロンプト：展開前）")
        st.code(user_view_before or "(empty)")
        st.caption("User（指示プロンプト：展開後＝実際に送信）")
        st.code(user_view_after or "(empty)")

        preview = [{"role": "system", "content": system_view}] + ag.history[-20:] + [{"role": "user", "content": user_view_after}]
        st.caption("送信 messages（展開後）")
        st.code(json_pretty(preview), language="json")

    st.session_state.grid[r][c] = ag


# ============ ルーティング ============
if st.session_state.view == "main":
    render_main()
else:
    render_detail()
